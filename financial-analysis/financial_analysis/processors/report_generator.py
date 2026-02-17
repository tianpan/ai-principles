"""
报告生成器模块

用于生成财务分析报告，包括文字内容和图表。
"""
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
import platform
import shutil
import tempfile
from pathlib import Path
import time
from matplotlib.font_manager import FontProperties
import uuid

# 设置非交互式后端
matplotlib.use('Agg')

# 配置中文字体
if platform.system() == "Windows":
    # Windows系统
    plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'SimSun']
elif platform.system() == "Darwin":
    # macOS系统
    plt.rcParams['font.sans-serif'] = ['PingFang SC', 'STHeiti', 'Heiti TC']
else:
    # Linux系统
    plt.rcParams['font.sans-serif'] = ['WenQuanYi Micro Hei', 'WenQuanYi Zen Hei', 'Droid Sans Fallback']

# 解决负号显示问题
plt.rcParams['axes.unicode_minus'] = False

class ReportGenerator:
    """财务报告生成器类"""
    
    def __init__(self, balance_sheet_data, income_statement_data, output_folder):
        """
        初始化报告生成器
        
        Args:
            balance_sheet_data: 资产负债表数据
            income_statement_data: 损益表数据
            output_folder: 输出文件夹路径
        """
        self.balance_sheet_data = balance_sheet_data
        self.income_statement_data = income_statement_data
        self.output_folder = output_folder
        
        # 使用临时目录来保存图表，避免权限问题
        self.temp_image_dir = tempfile.mkdtemp(prefix="financial_analysis_")
        
        # 确保输出目录存在
        self._ensure_output_folder()
    
    def _ensure_output_folder(self):
        """确保输出文件夹存在，如果不存在则创建"""
        try:
            os.makedirs(self.output_folder, exist_ok=True)
        except PermissionError:
            # 如果没有权限，使用用户临时目录
            self.output_folder = os.path.join(tempfile.gettempdir(), "financial_analysis_reports")
            os.makedirs(self.output_folder, exist_ok=True)
            print(f"权限不足，报告将保存到: {self.output_folder}")

    def _calculate_financial_ratios(self):
        """
        计算财务比率
        
        Returns:
            dict: 计算的财务比率
        """
        ratios = {
            'liquidity': {},      # 流动性比率
            'profitability': {},  # 盈利能力比率
            'operational': {},    # 运营效率比率
            'leverage': {}        # 杠杆比率
        }
        
        # 获取最新的日期
        if self.balance_sheet_data['dates']:
            latest_date = self.balance_sheet_data['dates'][-1]
        else:
            return ratios
            
        # 流动性比率计算
        # 流动比率 = 流动资产 / 流动负债
        current_assets = 0
        current_liabilities = 0
        
        for item, values in self.balance_sheet_data['assets'].items():
            if '流动资产' in item and latest_date in values:
                current_assets += values[latest_date] if not np.isnan(values[latest_date]) else 0
                
        for item, values in self.balance_sheet_data['liabilities'].items():
            if '流动负债' in item and latest_date in values:
                current_liabilities += values[latest_date] if not np.isnan(values[latest_date]) else 0
        
        if current_liabilities != 0:
            ratios['liquidity']['流动比率'] = current_assets / current_liabilities
        else:
            ratios['liquidity']['流动比率'] = np.nan
            
        # 盈利能力比率计算
        # 总资产收益率 (ROA) = 净利润 / 总资产
        net_profit = 0
        total_assets = 0
        total_revenue = 0
        
        for item, values in self.income_statement_data['profit'].items():
            if '净利润' in item and latest_date in values:
                net_profit += values[latest_date] if not np.isnan(values[latest_date]) else 0
                
        for item, values in self.balance_sheet_data['assets'].items():
            if '资产总计' in item and latest_date in values:
                total_assets += values[latest_date] if not np.isnan(values[latest_date]) else 0
                
        for item, values in self.income_statement_data['revenue'].items():
            if '营业收入' in item and latest_date in values:
                total_revenue += values[latest_date] if not np.isnan(values[latest_date]) else 0
                
        if total_assets != 0:
            ratios['profitability']['总资产收益率(ROA)'] = net_profit / total_assets
        else:
            ratios['profitability']['总资产收益率(ROA)'] = np.nan
            
        # 净利润率 = 净利润 / 营业收入
        if total_revenue != 0:
            ratios['profitability']['净利润率'] = net_profit / total_revenue
        else:
            ratios['profitability']['净利润率'] = np.nan
            
        # 杠杆比率计算
        # 资产负债率 = 总负债 / 总资产
        total_liabilities = 0
        
        for item, values in self.balance_sheet_data['liabilities'].items():
            if '负债合计' in item and latest_date in values:
                total_liabilities += values[latest_date] if not np.isnan(values[latest_date]) else 0
                
        if total_assets != 0:
            ratios['leverage']['资产负债率'] = total_liabilities / total_assets
        else:
            ratios['leverage']['资产负债率'] = np.nan
            
        return ratios
    
    def _generate_trend_data(self):
        """
        生成趋势分析数据
        
        Returns:
            dict: 趋势数据
        """
        trends = {
            'assets': {},
            'liabilities': {},
            'revenue': {},
            'profit': {}
        }
        
        # 处理资产趋势
        for item, values in self.balance_sheet_data['assets'].items():
            if '资产总计' in item or '流动资产' in item or '非流动资产' in item:
                trends['assets'][item] = values
                
        # 处理负债趋势
        for item, values in self.balance_sheet_data['liabilities'].items():
            if '负债合计' in item or '流动负债' in item or '非流动负债' in item:
                trends['liabilities'][item] = values
                
        # 处理收入趋势
        for item, values in self.income_statement_data['revenue'].items():
            if '营业收入' in item or '主营业务收入' in item:
                trends['revenue'][item] = values
                
        # 处理利润趋势
        for item, values in self.income_statement_data['profit'].items():
            if '净利润' in item or '营业利润' in item or '利润总额' in item:
                trends['profit'][item] = values
                
        return trends
    
    def _generate_charts(self):
        """
        生成报表图表
        
        Returns:
            dict: 图表文件路径
        """
        charts = {}
        
        # 生成财务比率图表
        ratios = self._calculate_financial_ratios()
        
        # 设置全局字体属性，解决中文显示问题
        plt.rcParams['font.family'] = plt.rcParams['font.sans-serif'][0]
        
        # 生成流动性比率图表
        if ratios['liquidity']:
            fig, ax = plt.subplots(figsize=(8, 6))
            
            names = list(ratios['liquidity'].keys())
            values = list(ratios['liquidity'].values())
            
            ax.bar(names, values, color='skyblue')
            ax.set_title('流动性比率')
            ax.set_ylabel('比率值')
            plt.tight_layout()
            
            # 使用时间戳确保文件名唯一
            timestamp = int(time.time())
            chart_path = os.path.join(self.temp_image_dir, f'liquidity_ratios_{timestamp}.png')
            
            try:
                plt.savefig(chart_path)
                charts['liquidity_ratios'] = chart_path
            except Exception as e:
                print(f"保存图表时出错: {e}")
            finally:
                plt.close()
            
        # 生成盈利能力比率图表
        if ratios['profitability']:
            fig, ax = plt.subplots(figsize=(8, 6))
            
            names = list(ratios['profitability'].keys())
            values = list(ratios['profitability'].values())
            
            ax.bar(names, values, color='lightgreen')
            ax.set_title('盈利能力比率')
            ax.set_ylabel('比率值')
            plt.tight_layout()
            
            # 使用时间戳确保文件名唯一
            timestamp = int(time.time())
            chart_path = os.path.join(self.temp_image_dir, f'profitability_ratios_{timestamp}.png')
            
            try:
                plt.savefig(chart_path)
                charts['profitability_ratios'] = chart_path
            except Exception as e:
                print(f"保存图表时出错: {e}")
            finally:
                plt.close()
            
        # 生成杠杆比率图表
        if ratios['leverage']:
            fig, ax = plt.subplots(figsize=(8, 6))
            
            names = list(ratios['leverage'].keys())
            values = list(ratios['leverage'].values())
            
            ax.bar(names, values, color='coral')
            ax.set_title('杠杆比率')
            ax.set_ylabel('比率值')
            plt.tight_layout()
            
            # 使用时间戳确保文件名唯一
            timestamp = int(time.time())
            chart_path = os.path.join(self.temp_image_dir, f'leverage_ratios_{timestamp}.png')
            
            try:
                plt.savefig(chart_path)
                charts['leverage_ratios'] = chart_path
            except Exception as e:
                print(f"保存图表时出错: {e}")
            finally:
                plt.close()
            
        # 生成趋势分析图表
        trends = self._generate_trend_data()
        
        # 生成资产趋势图表
        if trends['assets'] and self.balance_sheet_data['dates']:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            for item, values in trends['assets'].items():
                data_points = []
                for date in self.balance_sheet_data['dates']:
                    if date in values:
                        data_points.append(values[date])
                    else:
                        data_points.append(np.nan)
                
                ax.plot(self.balance_sheet_data['dates'], data_points, marker='o', label=item)
            
            ax.set_title('资产趋势分析')
            ax.set_xlabel('日期')
            ax.set_ylabel('金额')
            ax.legend()
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            chart_path = os.path.join(self.temp_image_dir, 'assets_trend.png')
            plt.savefig(chart_path)
            plt.close()
            
            charts['assets_trend'] = chart_path
            
        # 生成收入和利润趋势图表
        if (trends['revenue'] or trends['profit']) and self.income_statement_data['dates']:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            for item, values in trends['revenue'].items():
                data_points = []
                for date in self.income_statement_data['dates']:
                    if date in values:
                        data_points.append(values[date])
                    else:
                        data_points.append(np.nan)
                
                ax.plot(self.income_statement_data['dates'], data_points, marker='o', label=item)
                
            for item, values in trends['profit'].items():
                data_points = []
                for date in self.income_statement_data['dates']:
                    if date in values:
                        data_points.append(values[date])
                    else:
                        data_points.append(np.nan)
                
                ax.plot(self.income_statement_data['dates'], data_points, marker='s', label=item)
            
            ax.set_title('收入和利润趋势分析')
            ax.set_xlabel('日期')
            ax.set_ylabel('金额')
            ax.legend()
            plt.xticks(rotation=45)
            plt.tight_layout()
            
            chart_path = os.path.join(self.temp_image_dir, 'revenue_profit_trend.png')
            plt.savefig(chart_path)
            plt.close()
            
            charts['revenue_profit_trend'] = chart_path
            
        return charts
    
    def generate(self):
        """
        生成财务分析报告
        
        Returns:
            str: 生成的报告文件路径
        """
        # 生成唯一的文件名
        try:
            company_name = self.balance_sheet_data.get('company_name', "未知企业")
            if not isinstance(company_name, str) or not company_name:
                company_name = "未知企业"
            # 移除公司名称中的非法字符
            company_name = ''.join(c for c in company_name if c.isalnum() or c in (' ', '_', '-'))
            
            timestamp = int(time.time())
            output_filename = f"财务分析报告_{company_name}_{timestamp}.docx"
            output_path = os.path.join(self.output_folder, output_filename)
        
            # 尝试写入测试文件，验证权限
            test_file_path = os.path.join(self.output_folder, f"test_write_{timestamp}.tmp")
            try:
                with open(test_file_path, 'w') as f:
                    f.write("test")
                os.remove(test_file_path)
            except (PermissionError, OSError):
                # 如果测试写入失败，直接转到临时目录
                tmp_dir = tempfile.gettempdir()
                self.output_folder = tmp_dir
                output_path = os.path.join(tmp_dir, output_filename)
                print(f"输出目录无写入权限，切换到临时目录: {tmp_dir}")
            
            # 检查文件是否已存在
            if os.path.exists(output_path):
                # 如果文件存在，使用新的文件名
                base, ext = os.path.splitext(output_filename)
                output_filename = f"{base}_{int(time.time())}{ext}"
                output_path = os.path.join(self.output_folder, output_filename)
            
            # 生成报告
            success_path = self._generate_report(output_path)
            
            # 验证文件是否真的生成了
            if os.path.exists(success_path):
                return success_path
            else:
                # 如果文件不存在，说明生成失败，尝试临时目录
                tmp_dir = tempfile.gettempdir()
                output_path = os.path.join(tmp_dir, output_filename)
                return self._generate_report(output_path)
                
        except PermissionError as e:
            # 如果遇到权限错误，尝试写入临时目录
            tmp_dir = tempfile.gettempdir()
            output_filename = f"财务分析报告_{int(time.time())}.docx"
            output_path = os.path.join(tmp_dir, output_filename)
            print(f"原始路径无写入权限: {e}，尝试写入临时目录: {output_path}")
            return self._generate_report(output_path)
        except Exception as e:
            # 捕获所有其他异常，尝试写入临时目录
            print(f"报告生成过程中发生错误: {e}")
            tmp_dir = tempfile.gettempdir()
            output_filename = f"财务分析报告_错误恢复_{int(time.time())}.docx"
            output_path = os.path.join(tmp_dir, output_filename)
            return self._generate_report(output_path)
    
    def _generate_report(self, output_path):
        """
        生成财务分析报告的内部实现
        
        Args:
            output_path: 报告输出路径
            
        Returns:
            str: 实际保存的报告路径
        """
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('财务分析报告', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加报告生成日期
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(f"生成日期: {pd.Timestamp.now().strftime('%Y-%m-%d')}")
        date_run.font.size = Pt(10)
        
        # 添加企业基本信息
        doc.add_heading('1. 企业基本信息', level=1)
        doc.add_paragraph('本报告基于企业提供的财务数据进行分析，旨在提供企业财务状况的概述和评估。')
        
        # 添加财务状况概述
        doc.add_heading('2. 财务状况概述', level=1)
        
        # 资产负债情况
        doc.add_heading('2.1 资产负债情况', level=2)
        assets_liabilities_para = doc.add_paragraph()
        
        # 获取资产和负债总额
        total_assets = 0
        total_liabilities = 0
        total_equity = 0
        
        if self.balance_sheet_data.get('dates', []):
            latest_date = self.balance_sheet_data['dates'][-1]
            
            for item, values in self.balance_sheet_data.get('assets', {}).items():
                if '资产总计' in item and latest_date in values:
                    total_assets = values[latest_date]
                    
            for item, values in self.balance_sheet_data.get('liabilities', {}).items():
                if '负债合计' in item and latest_date in values:
                    total_liabilities = values[latest_date]
                    
            for item, values in self.balance_sheet_data.get('equity', {}).items():
                if '所有者权益' in item and latest_date in values:
                    total_equity = values[latest_date]
        
        assets_liabilities_para.add_run(f"截至最新报表日期，企业总资产为 {total_assets:,.2f}，"
                                       f"总负债为 {total_liabilities:,.2f}，"
                                       f"所有者权益为 {total_equity:,.2f}。")
        
        # 经营状况
        doc.add_heading('2.2 经营状况', level=2)
        operations_para = doc.add_paragraph()
        
        # 获取收入和利润数据
        total_revenue = 0
        net_profit = 0
        
        if self.income_statement_data.get('dates', []):
            latest_date = self.income_statement_data['dates'][-1]
            
            for item, values in self.income_statement_data.get('revenue', {}).items():
                if '营业收入' in item and latest_date in values:
                    total_revenue = values[latest_date]
                    
            for item, values in self.income_statement_data.get('profit', {}).items():
                if '净利润' in item and latest_date in values:
                    net_profit = values[latest_date]
        
        operations_para.add_run(f"最近一期企业实现营业收入 {total_revenue:,.2f}，"
                               f"净利润 {net_profit:,.2f}。")
        
        # 财务指标分析
        doc.add_heading('3. 财务指标分析', level=1)
        
        # 计算财务比率
        ratios = self._calculate_financial_ratios()
        
        # 添加流动性指标
        doc.add_heading('3.1 流动性指标', level=2)
        liquidity_para = doc.add_paragraph()
        if 'liquidity' in ratios and '流动比率' in ratios['liquidity']:
            liquidity_para.add_run(f"流动比率: {ratios['liquidity']['流动比率']:.2f}\n")
            liquidity_para.add_run("流动比率反映企业短期偿债能力，一般认为该比率以2:1为佳。")
        
        # 添加盈利能力指标
        doc.add_heading('3.2 盈利能力指标', level=2)
        profitability_para = doc.add_paragraph()
        if 'profitability' in ratios:
            if '总资产收益率(ROA)' in ratios['profitability']:
                profitability_para.add_run(f"总资产收益率(ROA): {ratios['profitability']['总资产收益率(ROA)']:.2%}\n")
            if '净利润率' in ratios['profitability']:
                profitability_para.add_run(f"净利润率: {ratios['profitability']['净利润率']:.2%}\n")
            profitability_para.add_run("以上指标反映企业的获利能力，数值越高表明盈利能力越强。")
        
        # 添加杠杆指标
        doc.add_heading('3.3 杠杆指标', level=2)
        leverage_para = doc.add_paragraph()
        if 'leverage' in ratios and '资产负债率' in ratios['leverage']:
            leverage_para.add_run(f"资产负债率: {ratios['leverage']['资产负债率']:.2%}\n")
            leverage_para.add_run("资产负债率反映企业的财务风险，一般认为该比率不超过70%为宜。")
        
        # 趋势分析
        doc.add_heading('4. 趋势分析', level=1)
        
        # 生成图表
        charts = self._generate_charts()
        
        # 添加资产负债趋势
        doc.add_heading('4.1 资产负债趋势', level=2)
        doc.add_paragraph('以下图表展示了企业资产的变化趋势:')
        if 'assets_trend' in charts:
            try:
                doc.add_picture(charts['assets_trend'], width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"图表添加失败: {e}")
        
        # 添加收入利润趋势
        doc.add_heading('4.2 收入利润趋势', level=2)
        doc.add_paragraph('以下图表展示了企业收入和利润的变化趋势:')
        if 'revenue_profit_trend' in charts:
            try:
                doc.add_picture(charts['revenue_profit_trend'], width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"图表添加失败: {e}")
        
        # 添加财务比率图表
        doc.add_heading('4.3 财务比率趋势', level=2)
        doc.add_paragraph('以下图表展示了企业主要财务比率:')
        
        if 'liquidity_ratios' in charts:
            try:
                doc.add_picture(charts['liquidity_ratios'], width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"流动性比率图表添加失败: {e}")
            
        if 'profitability_ratios' in charts:
            try:
                doc.add_picture(charts['profitability_ratios'], width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"盈利能力比率图表添加失败: {e}")
            
        if 'leverage_ratios' in charts:
            try:
                doc.add_picture(charts['leverage_ratios'], width=Inches(5))
            except Exception as e:
                doc.add_paragraph(f"杠杆比率图表添加失败: {e}")
        
        # 结论与建议
        doc.add_heading('5. 结论与建议', level=1)
        
        conclusion_para = doc.add_paragraph()
        
        # 基于财务比率给出基本评估
        if 'liquidity' in ratios and '流动比率' in ratios['liquidity']:
            if ratios['liquidity']['流动比率'] < 1:
                conclusion_para.add_run("企业短期偿债能力存在风险，建议加强营运资金管理。\n")
            elif ratios['liquidity']['流动比率'] > 2:
                conclusion_para.add_run("企业短期偿债能力较强，但可能存在资金利用效率有待提高的问题。\n")
            else:
                conclusion_para.add_run("企业短期偿债能力处于合理水平。\n")
                
        if 'leverage' in ratios and '资产负债率' in ratios['leverage']:
            if ratios['leverage']['资产负债率'] > 0.7:
                conclusion_para.add_run("企业负债水平较高，财务风险需要关注。\n")
            else:
                conclusion_para.add_run("企业资本结构较为合理，财务风险可控。\n")
                
        if 'profitability' in ratios and '净利润率' in ratios['profitability']:
            if ratios['profitability']['净利润率'] < 0:
                conclusion_para.add_run("企业处于亏损状态，需要关注成本控制和业务调整。\n")
            elif ratios['profitability']['净利润率'] < 0.05:
                conclusion_para.add_run("企业盈利能力偏弱，建议分析原因并采取改进措施。\n")
            else:
                conclusion_para.add_run("企业盈利能力良好，建议保持现有业务模式并寻求扩张机会。\n")
                
        # 添加免责声明
        doc.add_heading('免责声明', level=1)
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.add_run("本报告基于提供的财务数据生成，仅供参考，不构成投资建议。"
                               "分析结果的准确性取决于输入数据的质量和完整性。")
        
        # 尝试保存文档，添加异常处理
        try:
            # 确保目录存在
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # 尝试保存文件
            doc.save(output_path)
            print(f"报告已成功保存到：{output_path}")
            
            # 验证文件是否实际创建
            if not os.path.exists(output_path):
                raise FileNotFoundError(f"文件保存失败：{output_path}")
                
            return output_path
            
        except (PermissionError, OSError) as e:
            # 如果遇到权限错误或其他OS错误，尝试写入临时目录
            print(f"保存到原路径失败：{e}")
            tmp_dir = tempfile.gettempdir()
            file_name = os.path.basename(output_path)
            tmp_path = os.path.join(tmp_dir, file_name)
            
            try:
                doc.save(tmp_path)
                print(f"报告已保存到临时位置：{tmp_path}")
                return tmp_path
            except Exception as e2:
                # 最后尝试使用唯一的临时文件名
                print(f"保存到临时路径失败：{e2}")
                unique_path = os.path.join(tmp_dir, f"财务报告_{uuid.uuid4().hex}.docx")
                doc.save(unique_path)
                return unique_path
                
        except Exception as e:
            # 捕获所有其他异常
            print(f"保存文档时出现未知错误：{e}")
            
            # 最终尝试 - 使用Python的tempfile模块创建临时文件
            fd, tmp_path = tempfile.mkstemp(suffix='.docx', prefix='财务报告_')
            os.close(fd)  # 关闭文件描述符
            
            doc.save(tmp_path)
            return tmp_path 